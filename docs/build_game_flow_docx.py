from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(__file__).with_name("Huong_dan_luong_chuc_nang_game.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
CODE_FILL = "F5F7FA"
HEADER_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)
                set_cell_margins(row.cells[index])
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.05
    code_style.paragraph_format.left_indent = Inches(0.15)
    code_style.paragraph_format.right_indent = Inches(0.05)

    if "Lead" not in styles:
        lead = styles.add_style("Lead", 1)
    else:
        lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    lead.font.size = Pt(11)
    lead.font.color.rgb = MUTED
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.25


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Hướng Dẫn Học Luồng Chức Năng Trong Game")
    set_run_font(run, "Calibri", 24, RGBColor(11, 37, 69), True)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph(style="Lead")
    subtitle.add_run(
        "Tài liệu giải thích từng chức năng theo đúng code hiện có của project Unity "
        "3D Action Game: từ state machine, combat, spawn wave, UI cho đến online/Fusion."
    )

    meta = doc.add_paragraph(style="Lead")
    meta.add_run("Cách học đề xuất: ").bold = True
    meta.add_run(
        "mở script được nhắc tới, đọc luồng chạy, đặt breakpoint ở các hàm chính, "
        "rồi chạy Play Mode để quan sát dữ liệu thay đổi."
    )


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code Block")
        set_paragraph_shading(p, CODE_FILL)
        run = p.add_run(line.rstrip())
        set_run_font(run, "Consolas", 9)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_overview_table(doc):
    doc.add_heading("0. Bản đồ tổng quan", level=1)
    doc.add_paragraph(
        "Project hiện có 5 cụm chính. Khi đọc code, bạn nên xác định mình đang đứng ở cụm nào trước, "
        "sau đó lần theo hàm gọi và event/RPC liên quan."
    )

    rows = [
        ("Điều phối trận đấu", "Quản lý Loading, StartMatch, Pause, Victory, Lose", "GameManager.cs"),
        ("Nhân vật", "State machine Idle, Run, Attack, Slide, Hurt, Dead", "Character.cs, Player.cs, Enemy.cs"),
        ("Combat", "Bật/tắt vùng gây sát thương, trừ máu, chết, drop item", "DamageCaster.cs, Health.cs, DamageOrb.cs"),
        ("Màn chơi", "Spawn enemy, mở cổng, xác định thắng/thua", "Spawner.cs, SpawnPoint.cs, Gate.cs"),
        ("UI và online", "Panel UI, Supabase room, Photon Fusion avatar", "UIManager.cs, PanelGamePlay.cs, RoomService.cs, FusionPlayerAvatar.cs"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Cụm", "Vai trò", "File nên đọc"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], HEADER_FILL)
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_run_font(r, "Calibri", 10, RGBColor(0, 0, 0), True)
    for group, role, files in rows:
        cells = table.add_row().cells
        cells[0].text = group
        cells[1].text = role
        cells[2].text = files
    set_table_width(table, [1.5, 2.75, 2.25])

    doc.add_heading("Các enum nên nắm trước", level=2)
    add_code(
        doc,
        """
public enum GameState { Loading, Home, StartMatch, Pause, EndMatch, Victory, Lose }
public enum CharacterState { Idle, Run, Attack, Slide, Hurt, Dead }
public enum PickUpType { Health, Coin }
""",
    )


def add_section(doc, title, idea, flow, code_blocks=None, bullets=None):
    doc.add_heading(title, level=1)
    doc.add_heading("Ý tưởng", level=2)
    doc.add_paragraph(idea)
    if bullets:
        add_bullets(doc, bullets)
    doc.add_heading("Luồng chạy", level=2)
    add_numbers(doc, flow)
    if code_blocks:
        doc.add_heading("Code minh họa", level=2)
        for code in code_blocks:
            add_code(doc, code)


def build_doc():
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_overview_table(doc)

    add_section(
        doc,
        "1. Luồng vào game và trạng thái trận đấu",
        "`GameManager` là trung tâm điều phối trận đấu offline. Nó giữ `CurrentState`, bật/tắt `Time.timeScale`, mở UI gameplay, hiển thị UI thắng/thua và nghe sự kiện từ player/spawner.",
        [
            "`Awake()` tạo singleton `GameManager.Instance`.",
            "Nếu `autoFindSceneObjects = true`, `CacheSceneObjects()` tìm `Player` và toàn bộ `Spawner` trong scene.",
            "`Start()` gọi `SubscribeSceneEvents()` để nghe `player.Died` và `spawner.Cleared`.",
            "`ChangeState(initialState)` đưa game vào trạng thái đầu tiên, thường là `Loading`.",
            "`LoadingRoutine()` đợi `loadingDuration`, sau đó chuyển sang `stateAfterLoading`, thường là `StartMatch`.",
            "Khi player chết offline, `OnPlayerDied()` đợi `loseDelayAfterPlayerDeath` rồi gọi `Lose()`.",
            "Khi tất cả spawner đã clear, `OnSpawnerCleared()` gọi `Victory()`.",
        ],
        [
            """
private void Start()
{
    SubscribeSceneEvents();
    ChangeState(initialState);
}

private IEnumerator LoadingRoutine()
{
    yield return new WaitForSecondsRealtime(loadingDuration);
    ChangeState(stateAfterLoading);
}
""",
            """
case GameState.StartMatch:
    Time.timeScale = 1f;
    OpenGameplayUI();
    break;

case GameState.Lose:
    Time.timeScale = 0f;
    CloseGameplayUI();
    OpenGameOverUI();
    break;
""",
        ],
    )

    add_section(
        doc,
        "2. Luồng state machine của nhân vật",
        "`Character` là lớp cha cho `Player` và `Enemy`. Mỗi nhân vật đều chạy qua các state `Idle`, `Run`, `Attack`, `Slide`, `Hurt`, `Dead`. Lớp cha xử lý phần chung; lớp con chỉ cung cấp hướng di chuyển và hành vi riêng từng state.",
        [
            "`Update()` hoặc `FixedUpdate()` gọi `MoveCharacter(deltaTime)`.",
            "Nếu đang `Dead`, nhân vật dừng movement, dừng impact và set animator speed về 0.",
            "Nếu đang spawn dissolve, nhân vật bị khóa movement cho tới khi hiệu ứng kết thúc.",
            "`UpdateState(CurrentState, deltaTime)` cho state hiện tại tự xử lý logic riêng.",
            "Nếu state cho phép di chuyển, `GetMoveDirection()` trả về hướng mục tiêu.",
            "Hướng di chuyển được làm mượt bằng `Vector3.MoveTowards`.",
            "Animator được cập nhật qua `Speed` và `IsGrounded`.",
            "Nhân vật di chuyển bằng `CharacterController`, `Rigidbody` hoặc `transform.position`.",
        ],
        [
            """
protected abstract Vector3 GetMoveDirection();

public void SwitchToState(CharacterState newState, bool forceRestart = false)
{
    if (CurrentState == newState && !forceRestart) return;

    ExitState(CurrentState);
    CurrentState = newState;
    EnterState(CurrentState);
}
""",
            """
protected virtual void OnEnterAttack() { }
protected virtual void OnUpdateAttack(float deltaTime) { }
protected virtual void OnExitAttack() { }
""",
        ],
    )

    add_section(
        doc,
        "3. Luồng di chuyển của Player",
        "`Player.GetMoveDirection()` là nơi gom input từ bàn phím và joystick. Nếu có joystick và không có keyboard thì dùng joystick; nếu có keyboard thì input được làm mượt để nhân vật không giật.",
        [
            "Đọc keyboard bằng `ReadKeyboardMoveInput()`.",
            "Đọc joystick bằng `JoyStickController.direct`.",
            "Nếu chỉ joystick có input, trả về hướng joystick đã clamp về độ dài tối đa 1.",
            "Nếu keyboard có input, dùng `keyboardInputAcceleration` để tăng dần vector di chuyển.",
            "Khi nhả phím, dùng `keyboardInputDeceleration` để giảm dần về `Vector3.zero`.",
            "`Character.MoveCharacter()` nhận hướng này để di chuyển và xoay nhân vật.",
        ],
        [
            """
protected override Vector3 GetMoveDirection()
{
    Vector3 keyboardDirection = ReadKeyboardMoveInput();
    Vector3 joystickDirection = JoyStickController.direct;

    bool hasKeyboardInput = keyboardDirection.sqrMagnitude > 0.001f;
    bool hasJoystickInput = joystickDirection.sqrMagnitude > 0.001f;

    if (hasJoystickInput && !hasKeyboardInput)
        return Vector3.ClampMagnitude(joystickDirection, 1f);

    smoothedKeyboardMoveInput = Vector3.MoveTowards(
        smoothedKeyboardMoveInput,
        hasKeyboardInput ? keyboardDirection : Vector3.zero,
        inputSmoothingSpeed * Time.deltaTime
    );

    return Vector3.ClampMagnitude(smoothedKeyboardMoveInput, 1f);
}
""",
        ],
        bullets=[
            "Offline: UI gọi trực tiếp `Player.Attack()` hoặc `Player.Slide()`.",
            "Online: UI gọi `FusionPlayerAvatar.RequestAttack()` hoặc `FusionPlayerAvatar.RequestSlide()` để đi qua Fusion.",
        ],
    )

    add_section(
        doc,
        "4. Luồng tấn công combo của Player",
        "Player có combo tối đa 3 đòn. Bấm attack khi đang attack sẽ queue đòn tiếp theo. Nếu người chơi di chuyển trong lúc attack, combo đang queue bị hủy để ưu tiên điều khiển.",
        [
            "UI hoặc phím Enter gọi `Player.Attack()`.",
            "Nếu đang `Slide`, `Hurt` hoặc `Dead`, hàm thoát ngay.",
            "Nếu đang `Attack`, `QueueNextCombo()` tăng số combo được yêu cầu.",
            "Nếu chưa attack, set `requestedComboCount = 1` rồi vào state `Attack`.",
            "`OnEnterAttack()` gọi `BeginAttackCombo(1)`.",
            "`BeginAttack()` set timer và trigger animator `Attack`.",
            "Khi animation kết thúc, animation event gọi `AttackAnimationEnds()`.",
            "`CompleteCurrentCombo()` quyết định đánh tiếp combo hay thoát về `Idle/Run`.",
        ],
        [
            """
public void Attack()
{
    if (CurrentState == CharacterState.Slide
        || CurrentState == CharacterState.Hurt
        || CurrentState == CharacterState.Dead)
        return;

    if (CurrentState == CharacterState.Attack)
    {
        QueueNextCombo();
        return;
    }

    requestedComboCount = 1;
    SwitchToState(CharacterState.Attack);
}
""",
            """
private void CompleteCurrentCombo()
{
    if (!cancelQueuedCombosForMove
        && currentComboIndex < requestedComboCount
        && currentComboIndex < MaxComboCount)
    {
        BeginAttackCombo(currentComboIndex + 1);
        return;
    }

    FinishAttack();
}
""",
        ],
    )

    add_section(
        doc,
        "5. Luồng gây sát thương cận chiến",
        "`DamageCaster` là collider trigger nằm ở vũ khí hoặc vùng chém. Collider mặc định bị tắt. Khi animation attack tới frame gây damage, animation event bật collider; hết frame gây damage thì tắt collider.",
        [
            "`DamageCaster.OnTriggerEnter()` hoặc `OnTriggerStay()` gọi `TryApplyDamage(other)`.",
            "Nếu chủ thể online không có quyền apply damage local, hàm thoát để tránh nhân đôi sát thương.",
            "Nếu target đã bị đánh trong cùng một cửa sổ damage, bỏ qua bằng `damageTargetIdSet`.",
            "Tìm `Character` hoặc avatar network trên collider bị chạm.",
            "Offline: gọi `targetCharacter.ApplyDamage(damage, attackerPosition)`.",
            "Online: gọi `RequestDamage(...)` trên avatar network để gửi RPC tới StateAuthority.",
            "Sau khi đánh trúng, phát slash VFX và đánh dấu target đã bị damage.",
        ],
        [
            """
private void TryApplyDamage(Collider other)
{
    Character targetCharacter = other.GetComponentInParent<Character>();

    if (targetCharacter == null || !targetCharacter.CompareTag(targetTag))
        return;

    if (HasDamagedTarget(targetCharacter))
        return;

    targetCharacter.ApplyDamage(damage, attackerPosition);
    PlayHitVFX(other);
    MarkDamagedTarget(targetCharacter);
}
""",
            """
public void ApplyDamage(int damage, Vector3 attackPos = new Vector3())
{
    if (Health == null || Health.IsDead) return;

    Health.ApplyDamage(damage);

    if (Health.IsDead)
    {
        SwitchToState(CharacterState.Dead);
        return;
    }

    PlayMaterialsBlink();
    SwitchToState(CharacterState.Hurt, true);
    AddImpact(attackPos, HurtImpactForce);
}
""",
        ],
    )

    add_section(
        doc,
        "6. Luồng slide / né đòn",
        "Slide là một state riêng của `Player`. Khi slide, player không được attack, không nhận input di chuyển bình thường, `DamageCaster` bị tắt và nhân vật lao theo `slideDirection`.",
        [
            "UI hoặc phím Space gọi `Player.Slide()`.",
            "Nếu đang `Attack`, `Slide`, `Hurt` hoặc `Dead`, không cho slide.",
            "`GetSlideDirection()` lấy hướng input hiện tại; nếu không có input thì dùng `transform.forward`.",
            "Vào state `Slide` bằng `SwitchToState(CharacterState.Slide, true)`.",
            "`OnEnterSlide()` set timer, tắt damage caster và trigger animator `Slide`.",
            "`OnUpdateSlide()` tự đẩy nhân vật theo hướng slide.",
            "Hết timer hoặc animation event gọi `SlideAnimationEnds()`, player trở về `Run` hoặc `Idle`.",
        ],
        [
            """
public void Slide()
{
    if (CurrentState == CharacterState.Attack
        || CurrentState == CharacterState.Slide
        || CurrentState == CharacterState.Hurt
        || CurrentState == CharacterState.Dead)
        return;

    slideDirection = GetSlideDirection();
    SwitchToState(CharacterState.Slide, true);
}

protected override void OnUpdateSlide(float deltaTime)
{
    Vector3 movement = slideDirection * (slideDistance / slideDuration) * deltaTime;
    MoveBy(movement);
    RotateTowards(slideDirection, deltaTime);
}
""",
        ],
    )

    add_section(
        doc,
        "7. Luồng AI Enemy",
        "`Enemy` kế thừa `Character`, nhưng input không đến từ người chơi. Thay vào đó, `Enemy.GetMoveDirection()` hỏi `NavMeshAgent` hướng nên đi để đuổi player gần nhất.",
        [
            "`Awake()` lấy `NavMeshAgent` và tắt `updatePosition/updateRotation` để code tự điều khiển transform.",
            "`RefreshClosestPlayerTarget()` tìm player gần nhất còn sống và có đường NavMesh hợp lệ.",
            "Nếu target nằm trong tầm đánh, enemy dừng path và trả về `Vector3.zero`.",
            "Nếu tới thời điểm repath, enemy gọi `agent.SetDestination(target.position)`.",
            "Hướng di chuyển được lấy từ `agent.steeringTarget - transform.position`.",
            "`ApplyEnemySeparation()` giúp enemy tách nhau, tránh dồn thành một cụm.",
            "`OnUpdateIdle()` và `OnUpdateRun()` kiểm tra tầm đánh để chuyển sang state `Attack`.",
        ],
        [
            """
protected override Vector3 GetMoveDirection()
{
    if (IsTargetInAttackRange())
    {
        StopAgentPath();
        return Vector3.zero;
    }

    if (Time.time >= nextRepathTime)
    {
        nextRepathTime = Time.time + repathInterval;
        agent.SetDestination(target.position);
    }

    Vector3 direction = agent.steeringTarget - transform.position;
    return ApplyEnemySeparation(direction.normalized);
}
""",
            """
private void TryEnterAttackState()
{
    if (IsTargetInAttackRange())
        SwitchToState(CharacterState.Attack);
}
""",
        ],
    )

    add_section(
        doc,
        "8. Luồng spawn enemy, clear wave và mở cổng",
        "`Spawner` là vùng trigger. Khi player đi vào, spawner sinh enemy từ các `SpawnPoint`. Mỗi enemy được đăng ký sự kiện `Died`; khi toàn bộ enemy chết, spawner clear và mở cổng.",
        [
            "Player chạm collider của `Spawner`.",
            "`OnTriggerEnter()` gọi `SpawnCharacters()`.",
            "Nếu không có `NetworkRunner`, spawner chạy offline bằng `SpawnOfflineCharacters()`.",
            "Spawner lặp qua từng `SpawnPoint` và instantiate enemy prefab.",
            "Mỗi enemy tăng `aliveEnemyCount` và đăng ký `spawnedCharacter.Died += OnSpawnedCharacterDied`.",
            "Khi enemy chết, `OnSpawnedCharacterDied()` giảm `aliveEnemyCount`.",
            "Nếu `aliveEnemyCount <= 0`, `ClearSpawner()` được gọi.",
            "`ClearSpawner()` mở gate, broadcast network nếu cần, rồi phát event `Cleared` cho `GameManager`.",
        ],
        [
            """
private void OnSpawnedCharacterDied(Character spawnedCharacter)
{
    spawnedCharacter.Died -= OnSpawnedCharacterDied;
    aliveEnemyCount = Mathf.Max(aliveEnemyCount - 1, 0);

    if (aliveEnemyCount <= 0)
        ClearSpawner();
}

private void ClearSpawner()
{
    hasCleared = true;
    OpenGate();
    Cleared?.Invoke(this);
}
""",
            """
public void OpenGate()
{
    if (isOpen) return;
    isOpen = true;
    openCoroutine = StartCoroutine(OpenGateAnimation());
}
""",
        ],
    )

    add_section(
        doc,
        "9. Luồng máu, chết, dissolve và drop item",
        "`Health` chỉ quản lý số máu và phát event `HealthChanged`. Logic chết, dissolve, drop item nằm trong `Character` để player và enemy dùng chung.",
        [
            "`Health.ApplyDamage()` trừ máu và clamp về tối thiểu 0.",
            "`HealthChanged` báo cho UI hoặc network health biết máu đã đổi.",
            "`Character.ApplyDamage()` kiểm tra nếu máu đã chết thì chuyển sang state `Dead`.",
            "`OnEnterDead()` lưu vị trí drop, phát event `Died`, trigger animator `Dead` và tắt `DamageCaster`.",
            "Nếu không bị suppress dissolve, character chạy `StartMaterialDissolve()`.",
            "Sau dissolve, `DropItem()` instantiate item offline hoặc `Runner.Spawn()` item network.",
        ],
        [
            """
public void ApplyDamage(int damage)
{
    if (IsDead || damage <= 0) return;

    currentHealth = Mathf.Max(currentHealth - damage, 0);
    NotifyHealthChanged();
}
""",
            """
protected virtual void OnEnterDead()
{
    deathDropPosition = transform.position;
    NotifyDied();
    SetAnimatorTrigger(DeadParameter);
    DisableDamageCaster();
    StartMaterialDissolve();
}
""",
        ],
    )

    add_section(
        doc,
        "10. Luồng nhặt item, hồi máu và cộng coin",
        "`PickUp` xử lý item rơi trong map. Offline thì apply trực tiếp vào `Character`; online thì gửi request/RPC để đồng bộ pickup cho các client.",
        [
            "Player chạm collider của pickup.",
            "`PickUp.OnTriggerEnter()` kiểm tra đã collected chưa.",
            "Nếu gặp `FusionPlayerAvatar`, đi theo luồng online `CollectNetworkPlayer()`.",
            "Nếu offline, kiểm tra tag `Player` rồi lấy `Character`.",
            "Gọi `character.ApplyPickupValue(type, value)`.",
            "`PickUpType.Health` gọi `AddHealth(value)`; `PickUpType.Coin` gọi `AddCoin(value)`.",
            "`AddCoin()` phát event `CoinChanged`.",
            "`PanelGamePlay.OnPlayerCoinChanged()` cập nhật text và lưu tổng coin lên Supabase nếu người chơi đã đăng nhập.",
        ],
        [
            """
private void OnTriggerEnter(Collider other)
{
    Character character = other.GetComponentInParent<Character>();
    collected = true;
    character.ApplyPickupValue(type, value);
    PlayCollectedVFX(transform.position);
    Destroy(gameObject);
}
""",
            """
public void ApplyPickupValue(PickUpType pickupType, int value)
{
    switch (pickupType)
    {
        case PickUpType.Health: AddHealth(value); break;
        case PickUpType.Coin: AddCoin(value); break;
    }
}
""",
        ],
    )

    add_section(
        doc,
        "11. Luồng UI gameplay",
        "`UIManager` load prefab UI từ `Resources/UI/` và mở panel theo type. `PanelGamePlay` là panel chính khi vào trận: attack, slide, health, coin, timer, FPS/ping và revive online.",
        [
            "`UIManager.Awake()` load toàn bộ `UICanvas` prefab trong `Resources/UI/`.",
            "`OpenUI<T>()` lấy hoặc instantiate panel, gọi `SetUp()`, rồi `Open()`.",
            "`PanelGamePlay.OnEnable()` bind button, tìm player local và subscribe health/coin.",
            "Attack button gọi `OnAttackButtonClicked()`.",
            "Slide button gọi `OnSlideButtonClicked()`.",
            "Nếu đang online, panel ưu tiên gọi `FusionPlayerAvatar`; nếu offline thì gọi `Player`.",
            "`OnPlayerHealthChanged()` cập nhật slider máu.",
            "`UpdateReviveButton()` tìm đồng đội downed gần nhất và xử lý giữ nút revive.",
        ],
        [
            """
public T OpenUI<T>() where T : UICanvas
{
    T canvas = GetUI<T>();
    canvas.SetUp();
    canvas.Open();
    return canvas;
}
""",
            """
public void OnAttackButtonClicked()
{
    if (fusionPlayerAvatar != null)
        fusionPlayerAvatar.RequestAttack();
    else if (player != null)
        player.Attack();
}
""",
        ],
    )

    add_section(
        doc,
        "12. Luồng online: đăng nhập, tạo phòng và vào trận",
        "Phần online chia làm hai lớp: Supabase quản lý tài khoản/phòng/match metadata; Photon Fusion quản lý gameplay realtime trong GameScene.",
        [
            "`AuthService.SignIn()` gọi Supabase Auth endpoint để lấy access token.",
            "Thông tin user, display name, avatar và coin được lưu vào `SupabaseSession`.",
            "`RoomService.CreateRoom()` gọi Edge Function `create_room`.",
            "`RoomService.JoinRoom()` gọi Edge Function `join_room`.",
            "`PanelRoomMatch` hiển thị danh sách người chơi và trạng thái ready.",
            "Host bấm Start Match, `StartMatchAsHostRoutine()` gọi `roomService.StartMatch()`.",
            "Khi Supabase có active match, `BeginLoadMatch()` lưu match vào `OnlineRoomSession` và load `GameScene`.",
        ],
        [
            """
SupabaseSession.AccessToken = response.access_token;
SupabaseSession.UserId = response.user.id;
SupabaseSession.Email = response.user.email;
SupabaseSession.DisplayName = response.user.GetDisplayName();
yield return LoadUserProfile(SupabaseSession.UserId, ...);
""",
            """
private void BeginLoadMatch(RoomService.MatchData match)
{
    OnlineRoomSession.SetMatch(match);
    OnlineRoomSession.CacheExpectedMatchPlayerCount();
    OnlineMatchLoadingOverlay.LoadScene(gameSceneName);
}
""",
        ],
    )

    add_section(
        doc,
        "13. Luồng Fusion khi vào GameScene online",
        "`FusionMatchBootstrap` là entry point của online match trong scene game. Nó tạo `NetworkRunner`, join session theo `MatchId`, rồi spawn player network local.",
        [
            "`Start()` lấy `OnlineRoomSession.MatchId` làm `SessionName`.",
            "Nếu không có match id và cho fallback, scene giữ player offline.",
            "`OnlineMatchStats.StartMatch()` bắt đầu ghi thống kê trận.",
            "Scene player template bị tắt để tránh trùng player offline và online.",
            "Script tạo `NetworkRunner` và `NetworkSceneManagerDefault`.",
            "`runner.StartGame()` join Photon Fusion session ở `GameMode.Shared`.",
            "Khi join thành công, `SpawnLocalPlayerIfNeeded()` spawn prefab network player ở spawn point tương ứng.",
            "`runner.SetPlayerObject()` gắn network object với `PlayerRef` local.",
        ],
        [
            """
StartGameResult result = await runner.StartGame(new StartGameArgs
{
    GameMode = GameMode.Shared,
    SessionName = sessionName,
    PlayerCount = maxPlayers,
    AuthValues = new AuthenticationValues(GetPhotonUserId())
});

SpawnLocalPlayerIfNeeded(runner.LocalPlayer);
""",
            """
localPlayerObject = runner.Spawn(networkPlayerPrefab, spawnPosition, spawnRotation, player);
runner.SetPlayerObject(player, localPlayerObject);
""",
        ],
    )

    add_section(
        doc,
        "14. Luồng FusionPlayerAvatar",
        "`FusionPlayerAvatar` thay thế `Player` offline khi chơi online. Nó xử lý authority, input, movement, attack, slide, damage, revive, nameplate và đồng bộ health.",
        [
            "`Spawned()` resolve reference, subscribe health, khởi tạo revive state và áp dụng authority.",
            "`ApplyAuthorityState()` phân biệt local/remote; local có tag `Player`, remote là `Untagged`.",
            "`player.enabled = false` để movement online không bị script `Player` offline tranh quyền.",
            "UI gọi `RequestAttack()` hoặc `RequestSlide()` để queue hành động.",
            "`FixedUpdateNetwork()` consume queue, cập nhật damage window và simulate movement.",
            "Attack online dùng controlled damage window để bật/tắt `DamageCaster` theo timer.",
            "Damage gửi qua `RPC_ApplyDamage()` tới StateAuthority.",
            "`FusionNetworkHealth` mirror local health lên network và apply network health về client render.",
        ],
        [
            """
public override void Spawned()
{
    ResolveReferences();
    SubscribeHealth();
    InitializeReviveState();
    ApplyAuthorityState();
    SetLocalIdentityIfNeeded();
    RefreshDisplayNameView();
}
""",
            """
public bool RequestDamage(int damage, Vector3 attackPosition, int damageSourceId = 0)
{
    RPC_ApplyDamage(damage, attackPosition, damageSourceId);
    return true;
}

[Rpc(RpcSources.All, RpcTargets.StateAuthority)]
private void RPC_ApplyDamage(int damage, Vector3 attackPosition, int damageSourceId)
{
    player.ApplyDamage(damage, attackPosition);
    networkHealth?.ForceSyncNow();
}
""",
        ],
    )

    add_section(
        doc,
        "15. Luồng chết, down và revive online",
        "Online không luôn xem player hết máu là loại khỏi trận ngay. Nếu còn lượt revive, player chuyển sang trạng thái downed; nếu hết lượt revive thì eliminated.",
        [
            "`HealthChanged` gọi `FusionPlayerAvatar.OnHealthChanged()`.",
            "Nếu health về 0 và chưa apply death, gọi `ApplyNetworkDeath()`.",
            "Nếu `RevivesRemaining > 0`, StateAuthority set `IsDowned = true`.",
            "Nếu hết revive, StateAuthority set `IsEliminated = true`.",
            "Character chuyển sang `Dead`; nếu downed thì suppress dissolve để còn revive được.",
            "Player còn sống đứng gần player downed, `PanelGamePlay` hiện nút revive.",
            "Giữ nút đủ `reviveHoldDuration`, panel gọi `RequestReviveTarget()`.",
            "Target nhận `RPC_RequestRevive()`, kiểm tra khoảng cách/trạng thái rồi hồi máu bằng `RPC_ApplyRevive()`.",
        ],
        [
            """
private void OnHealthChanged(int current, int max)
{
    if (current > 0 || hasAppliedNetworkDeath) return;
    ApplyNetworkDeath();
}
""",
            """
[Rpc(RpcSources.All, RpcTargets.StateAuthority)]
private void RPC_RequestRevive(PlayerRef reviver)
{
    if (!CanBeRevived) return;

    RevivesRemaining--;
    IsDowned = false;
    IsEliminated = false;
    health.SetHealthFromNetwork(reviveHealth, maxHealth);
    RPC_ApplyRevive(reviveHealth);
}
""",
        ],
    )

    add_section(
        doc,
        "16. Luồng thắng/thua online",
        "`NetworkMatchManager` đánh giá kết quả trận online. Mỗi client đều có thể quan sát trạng thái map, nhưng kết quả được broadcast qua `FusionPlayerAvatar` để mọi client áp dụng giống nhau.",
        [
            "`Update()` chỉ chạy đánh giá nếu online match đang active.",
            "Sau `evaluationStartDelay`, manager kiểm tra theo chu kỳ `evaluationInterval`.",
            "Nếu tất cả spawner đã clear, trận Victory.",
            "Nếu hết thời gian giới hạn, trận Lose.",
            "Nếu tất cả player không thể tiếp tục, trận Lose.",
            "`FinishMatch()` broadcast kết quả qua avatar local/fallback.",
            "Mọi client nhận `RPC_ApplyMatchResult()`.",
            "`NetworkMatchManager.ApplyNetworkResult()` gọi `GameManager.Victory()` hoặc `GameManager.Lose()`.",
        ],
        [
            """
private void EvaluateMatchState()
{
    if (AreAllSpawnersCleared())
    {
        FinishMatch(GameState.Victory);
        return;
    }

    if (HasMatchTimeExpired() || AreAllPlayersUnableToContinue())
        FinishMatch(GameState.Lose);
}
""",
            """
[Rpc(RpcSources.All, RpcTargets.All)]
private void RPC_ApplyMatchResult(int resultStateValue)
{
    GameState resultState = (GameState)resultStateValue;
    NetworkMatchManager.Ensure().ApplyNetworkResult(resultState);
}
""",
        ],
    )

    doc.add_heading("17. Thứ tự học và điểm đặt breakpoint", level=1)
    doc.add_paragraph(
        "Để học hiệu quả, đừng đọc toàn bộ project từ trên xuống. Hãy đi theo luồng hành vi, đặt breakpoint ở đầu luồng, rồi bước từng hàm."
    )
    add_numbers(
        doc,
        [
            "`GameManager.ChangeState()` để hiểu trận đấu bắt đầu/kết thúc.",
            "`Character.MoveCharacter()` để hiểu vòng lặp nhân vật.",
            "`Player.Attack()` và `Player.Slide()` để hiểu input offline.",
            "`DamageCaster.TryApplyDamage()` và `Character.ApplyDamage()` để hiểu combat.",
            "`Enemy.GetMoveDirection()` để hiểu AI.",
            "`Spawner.SpawnCharacters()` và `Spawner.ClearSpawner()` để hiểu wave/gate.",
            "`PanelGamePlay.OnAttackButtonClicked()` để hiểu UI nối vào gameplay.",
            "`RoomService.StartMatch()` và `FusionMatchBootstrap.Start()` để hiểu online match.",
            "`FusionPlayerAvatar.FixedUpdateNetwork()` để hiểu gameplay online.",
            "`NetworkMatchManager.EvaluateMatchState()` để hiểu thắng/thua online.",
        ],
    )
    doc.add_heading("Các event/RPC cần nhớ", level=2)
    add_bullets(
        doc,
        [
            "`Character.Died` được `GameManager` và `Spawner` nghe.",
            "`Health.HealthChanged` được `PanelGamePlay`, `FusionPlayerAvatar` và `FusionNetworkHealth` nghe.",
            "`Spawner.Cleared` được `GameManager` nghe offline và được broadcast online qua `FusionPlayerAvatar`.",
            "`RPC_ApplyDamage` chỉ chạy trên StateAuthority, sau đó health được mirror về các client.",
        ],
    )

    return doc


if __name__ == "__main__":
    document = build_doc()
    document.save(OUT_PATH)
    print(OUT_PATH)
