from docx import Document


PATH = r"C:\Users\hoang\Downloads\64HTTT1_2251161947_VUHOANGANH_rut_gon_sua_ma_hoa.docx"


NETWORK_PARAGRAPH = (
    "Kết quả kiểm thử lỗi mạng cho thấy hệ thống đã tách được dữ liệu giữa các phòng "
    "và trận thông qua room_id, match_id và phiên mạng của Photon Fusion. Trong quá trình "
    "chơi, nếu một người chơi thường hoặc chủ phòng thoát giữa chừng thì trận đấu của những "
    "người còn lại vẫn tiếp tục bình thường; trường hợp chỉ còn duy nhất một người chơi, trận "
    "vẫn hoạt động cho đến khi kết thúc. Tuy nhiên, các tình huống như kết nối lại sau mất mạng "
    "hoặc dọn phòng treo vẫn cần hoàn thiện thêm."
)

LIMIT_PARAGRAPH = (
    "Hạn chế hiện tại là quy mô kiểm thử còn nhỏ, chưa đánh giá tải lớn trong thời gian dài và "
    "chưa hoàn thiện các chức năng như kết nối lại sau mất mạng, chuyển chủ phòng trong giai đoạn "
    "lobby hoặc dọn phòng treo. Đây là các hướng cần tiếp tục tối ưu trong phiên bản sau."
)

ROW_PLAYER_LEAVE = [
    "Người chơi thường thoát giữa trận",
    "Số lượng người chơi trong trận giảm, có thể gây sai lệch đồng bộ hoặc làm trận dừng ngoài ý muốn",
    "Các người chơi còn lại vẫn tiếp tục trận bình thường; nếu chỉ còn 1 người chơi thì trận vẫn hoạt động đến khi kết thúc",
    "Đạt",
]

ROW_HOST_LEAVE = [
    "Chủ phòng thoát giữa trận",
    "Trận có thể bị hủy hoặc các người chơi còn lại bị ảnh hưởng do mất chủ phòng",
    "Trận vẫn tiếp tục với các người chơi còn lại; nếu chỉ còn 1 người chơi thì trận vẫn hoạt động bình thường đến khi kết thúc",
    "Đạt",
]


def set_row_text(row, values):
    for cell, value in zip(row.cells, values):
        cell.text = value


def main():
    doc = Document(PATH)

    doc.paragraphs[764].text = NETWORK_PARAGRAPH
    doc.paragraphs[768].text = LIMIT_PARAGRAPH

    table = doc.tables[40]
    first_col = [row.cells[0].text.strip() for row in table.rows]

    if ROW_PLAYER_LEAVE[0] not in first_col:
        row_a = table.add_row()
        set_row_text(row_a, ROW_PLAYER_LEAVE)

        row_b = table.add_row()
        set_row_text(row_b, ROW_HOST_LEAVE)

        target = table.rows[3]._tr
        target.addnext(row_b._tr)
        target.addnext(row_a._tr)
    else:
        for row in table.rows:
            key = row.cells[0].text.strip()
            if key == ROW_PLAYER_LEAVE[0]:
                set_row_text(row, ROW_PLAYER_LEAVE)
            if key == ROW_HOST_LEAVE[0]:
                set_row_text(row, ROW_HOST_LEAVE)

    doc.save(PATH)


if __name__ == "__main__":
    main()
