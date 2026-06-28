from docx import Document


INPUT_PATH = r"C:\Users\hoang\Downloads\64HTTT1_2251161947_VUHOANGANH_rut_gon.docx"
OUTPUT_PATH = r"C:\Users\hoang\Downloads\64HTTT1_2251161947_VUHOANGANH_rut_gon_sua_ma_hoa.docx"


REPLACEMENTS = {
    638: "Đồ án triển khai game bằng Unity, sử dụng Supabase để xác thực và lưu dữ liệu nghiệp vụ, đồng thời dùng Photon Fusion để đồng bộ trận đấu thời gian thực. Cách kết hợp này đáp ứng đồng thời yêu cầu đồ họa 3D, quản lý tài khoản, phòng chơi, trận đấu và quá trình chơi nhiều người.",
    640: "Phía ứng dụng người chơi chịu trách nhiệm hiển thị giao diện, điều khiển nhân vật, camera, hiệu ứng, HUD và kết nối tới các dịch vụ trực tuyến.",
    641: "Các công cụ phát triển phía ứng dụng được trình bày trong Bảng 4.1.",
    644: "Trong Unity, hệ thống được tổ chức thành HomeScene cho các chức năng ngoài trận và GameScene cho quá trình chơi chính.",
    645: "Mã nguồn được chia thành bốn nhóm chính: trực tuyến, quá trình chơi, giao diện và kiểu dữ liệu/trạng thái; cách tổ chức này giúp dễ quản lý, mở rộng và kiểm thử.",
    648: "Phía máy chủ sử dụng Supabase để cung cấp xác thực người chơi, cơ sở dữ liệu PostgreSQL, cơ chế thời gian thực và các hàm xử lý nghiệp vụ.",
    649: "Các công cụ phát triển phía máy chủ được trình bày trong Bảng 4.2.",
    653: "Supabase Auth được dùng để xác thực bằng email và mật khẩu, sau đó cấp mã phiên để client gọi các hàm phía máy chủ.",
    654: "Cơ sở dữ liệu lưu các bảng users, rooms, room_players, matches và match_players nhằm quản lý tài khoản, phòng chơi, trận đấu và thống kê người chơi.",
    655: "Các thao tác quan trọng như tạo phòng, tham gia phòng, bắt đầu trận và lưu kết quả đều đi qua Edge Functions để kiểm tra quyền và ràng buộc dữ liệu, qua đó hạn chế client ghi sai hoặc ghi vượt quyền.",
    657: "Trong trận đấu, hệ thống dùng Photon Fusion để đồng bộ dữ liệu thời gian thực giữa các client.",
    658: "Những trạng thái được đồng bộ gồm vị trí, hướng xoay, hành động, máu, trạng thái sống/chết, hồi sinh, kẻ địch, vật phẩm, mở cổng và kết quả trận.",
    659: "Sau khi trận được tạo, Unity sử dụng match_id để tham gia phiên mạng tương ứng; mỗi người chơi được sinh thành một network player và các hành động trong trận được truyền qua cơ chế đồng bộ của Photon Fusion.",
    660: "Việc tách vai trò giữa Supabase và Photon Fusion giúp dữ liệu nghiệp vụ được lưu bền vững, còn dữ liệu quá trình chơi được cập nhật nhanh theo thời gian thực.",
    663: "Mã nguồn được quản lý bằng Git và GitHub. Postman cùng Supabase Dashboard được sử dụng để kiểm thử function, theo dõi dữ liệu và kiểm tra trạng thái phòng, trận cũng như kết quả sau trận.",
    667: "Sau quá trình triển khai, hệ thống đã hoàn thành các chức năng chính của game 3D co-op nhiều người chơi, từ đăng nhập, quản lý phòng và trận đến đồng bộ quá trình chơi, AI kẻ địch, lưu kết quả và hiển thị lịch sử trận đấu.",
    669: "Hệ thống đã triển khai đăng ký, đăng nhập qua Supabase Auth, lưu thông tin phiên đăng nhập và tải hồ sơ người chơi từ máy chủ. Các thông tin như tên hiển thị, ảnh đại diện và coin được lưu tập trung nên không bị mất khi người chơi thoát game hoặc đăng nhập lại.",
    680: "Khi người chơi tạo phòng, phía máy chủ sinh mã phòng và lưu dữ liệu vào cơ sở dữ liệu; người chơi khác có thể nhập mã để tham gia. Trong phòng, hệ thống hiển thị danh sách thành viên, vai trò chủ phòng/người chơi và trạng thái sẵn sàng, đồng thời cập nhật thay đổi để mọi thành viên theo dõi cùng một trạng thái.",
    686: "Người chơi trong phòng có thể chuyển trạng thái sẵn sàng; chủ phòng chỉ được bắt đầu trận khi các điều kiện hợp lệ được đáp ứng. Khi trận bắt đầu, phía máy chủ tạo dữ liệu trận, gắn match_id riêng và chuyển các thành viên của phòng sang GameScene để vào cùng một phiên chơi.",
    692: "Sau khi trận được tạo, Unity dùng match_id để tham gia phiên Photon Fusion tương ứng. Hệ thống đồng bộ được vị trí, hướng xoay, hành động, máu, tên hiển thị và trạng thái người chơi, đồng thời tách biệt dữ liệu giữa các trận để người chơi ở phòng này không nhận dữ liệu của phòng khác.",
    696: "Trong quá trình chơi, người chơi có thể di chuyển, tấn công, trượt né, nhận sát thương và theo dõi máu qua HUD. Nhân vật được tổ chức theo các trạng thái như Idle, Run, Attack, Slide, Hurt và Dead để việc điều khiển và xử lý hành vi rõ ràng hơn.",
    704: "Game triển khai AI kẻ địch bằng Unity NavMesh; kẻ địch có thể tìm đường, chọn mục tiêu gần nhất, di chuyển và tấn công người chơi. Kẻ địch được sinh theo khu vực hoặc theo đợt; khi một khu vực được dọn sạch, cổng sẽ mở để người chơi sang khu vực tiếp theo và độ khó có thể tăng theo đợt hoặc theo số người chơi.",
    708: "Hệ thống đã triển khai chiến đấu giữa người chơi và kẻ địch, trong đó sát thương và trạng thái máu/chết được đồng bộ giữa các client. Ở chế độ co-op, người chơi có thể hồi sinh đồng đội trong phạm vi và thời gian cho phép, qua đó tăng tính phối hợp giữa các thành viên.",
    715: "Trong trận đấu, người chơi có thể nhặt vật phẩm như coin và health orb. Khi vật phẩm đã được nhặt, trạng thái của vật phẩm được đồng bộ để tránh bị nhặt lặp; coin của người chơi cũng được cập nhật trên giao diện và có thể lưu lại sau trận.",
    719: "Hệ thống kết thúc trận theo điều kiện thắng hoặc thua, sau đó tổng hợp thống kê như số kẻ địch tiêu diệt, lượng sát thương gây ra, số lần bị hạ, số lần hồi sinh và thời gian sống sót để lưu vào lịch sử trận.",
    720: "Giao diện kết quả hiển thị bảng xếp hạng của người chơi trong trận, đồng thời cho phép xem lại lịch sử các trận đã tham gia.",
    728: "Sau khi hoàn thành các chức năng chính, hệ thống được kiểm thử để đánh giá tính đúng chức năng, khả năng đồng bộ nhiều người chơi, độ ổn định phía máy chủ và mức độ an toàn dữ liệu.",
    730: "Kiểm thử tập trung vào bốn nhóm mục tiêu: xác nhận luồng chức năng chính, đánh giá đồng bộ trong trận, kiểm tra việc kiểm soát nghiệp vụ ở phía máy chủ và quan sát hệ thống trong phạm vi 2 đến 4 người chơi.",
    733: "Việc kiểm thử được thực hiện bằng cách cho nhiều client đăng nhập bằng các tài khoản khác nhau, cùng tạo hoặc tham gia phòng, vào trận và thực hiện các thao tác chính trong cùng một phiên chơi.",
    748: "Với game đa người chơi, yếu tố quan trọng nhất là khả năng giữ trạng thái nhất quán giữa các client trong cùng một trận.",
    751: "Kết quả đo cho thấy hệ thống hoạt động ổn định trong phạm vi thử nghiệm từ 2 đến 4 người chơi. Ở kịch bản cơ bản với 2 người chơi, FPS trung bình đạt 43 và RTT khoảng 71 ms.",
    753: "Khi tăng lên 4 người chơi, FPS duy trì quanh 40 và RTT tăng lên 87 đến 100 ms, kể cả trong kịch bản chiến đấu với nhiều kẻ địch. Ở bước kết thúc trận và lưu kết quả, FPS tăng lên 69 và RTT khoảng 65 ms do tải đồng bộ thời gian thực giảm.",
    754: "Nhìn chung, bản demo đáp ứng tốt quy mô thử nghiệm hiện tại, nhưng chưa đủ cơ sở để kết luận về tải lớn hoặc nhiều phòng hoạt động đồng thời trong thời gian dài.",
    756: "Các thao tác quan trọng như tạo phòng, tham gia phòng, bắt đầu trận và lưu kết quả đều được kiểm thử qua Supabase Edge Functions nhằm bảo đảm client không thể tự ý ghi dữ liệu không hợp lệ.",
    765: "Kết quả kiểm thử lỗi mạng cho thấy hệ thống đã tách được dữ liệu giữa các phòng và trận thông qua room_id, match_id và phiên mạng của Photon Fusion. Tuy nhiên, các tình huống như chủ phòng thoát giữa chừng, kết nối lại sau mất mạng hoặc dọn phòng treo vẫn cần hoàn thiện thêm.",
    767: "Kết quả kiểm thử cho thấy hệ thống đáp ứng được luồng chính của đồ án: người chơi có thể đăng nhập, tạo phòng, tham gia phòng, sẵn sàng, bắt đầu trận, chơi cùng nhau, đồng bộ trạng thái trong trận và lưu kết quả sau trận.",
    768: "Supabase đảm nhiệm dữ liệu nghiệp vụ cần lưu lâu dài và kiểm soát quyền, còn Photon Fusion đảm nhiệm dữ liệu thay đổi liên tục trong trận như vị trí, hành động, máu, kẻ địch, vật phẩm và trạng thái thắng/thua. Cách phân chia này phù hợp với phạm vi game co-op 2 đến 4 người chơi.",
    769: "Hạn chế hiện tại là quy mô kiểm thử còn nhỏ, chưa đánh giá tải lớn trong thời gian dài và chưa hoàn thiện các chức năng như kết nối lại, chuyển chủ phòng hay dọn phòng treo. Đây là các hướng cần tiếp tục tối ưu trong phiên bản sau.",
}


def main():
    doc = Document(INPUT_PATH)
    for idx, text in REPLACEMENTS.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
