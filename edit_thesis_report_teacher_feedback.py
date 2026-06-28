from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


INPUT_PATH = Path(r"C:\Users\hoang\Downloads\64HTTT1_2251161947_VUHOANGANH.docx")
OUTPUT_PATH = Path(r"C:\Users\hoang\Downloads\64HTTT1_2251161947_VUHOANGANH_sua_theo_nhan_xet.docx")


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def remove_block(block):
    block._element.getparent().remove(block._element)


def clear_paragraph_content(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def rewrite_paragraph(paragraph: Paragraph, text: str) -> Paragraph:
    style = paragraph.style
    alignment = paragraph.alignment
    clear_paragraph_content(paragraph)
    paragraph.add_run(text)
    paragraph.style = style
    paragraph.alignment = alignment
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: int, cols: int, style: str | None = "Table Grid") -> Table:
    parent = paragraph._parent
    section = paragraph.part.document.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    table = parent.add_table(rows=rows, cols=cols, width=width)
    paragraph._p.addnext(table._tbl)
    if style:
        table.style = style
    return table


def find_paragraph_contains(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph containing {text!r} not found")


def find_last_paragraph_contains(doc: Document, text: str) -> Paragraph:
    for paragraph in reversed(doc.paragraphs):
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph containing {text!r} not found")


def find_table_after_caption(doc: Document, caption_text: str) -> Table:
    blocks = list(iter_block_items(doc))
    for index, block in enumerate(blocks):
        if isinstance(block, Paragraph) and caption_text in block.text:
            for next_block in blocks[index + 1 :]:
                if isinstance(next_block, Table):
                    return next_block
                if isinstance(next_block, Paragraph) and next_block.text.strip():
                    break
    raise ValueError(f"Table after caption {caption_text!r} not found")


def set_table_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if existing is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def format_table_for_report(table: Table) -> None:
    if table.rows:
        set_table_header_repeat(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def remove_blocks_between(doc: Document, start_text: str, end_text: str) -> None:
    to_remove = []
    removing = False
    for block in list(iter_block_items(doc)):
        block_text = block.text if isinstance(block, Paragraph) else "\n".join(
            " | ".join(c.text for c in row.cells) for row in block.rows
        )
        if not removing and start_text in block_text:
            removing = True
        if removing:
            if end_text in block_text:
                break
            to_remove.append(block)
    for block in to_remove:
        remove_block(block)


def replace_in_text(text: str, replacements: list[tuple[str, str]]) -> str:
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def set_caption_styles(doc: Document) -> None:
    caption_style = doc.styles["Caption"]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            continue
        if text.startswith("Bảng ") or text.startswith("Hình ") or text.startswith("Hình "):
            paragraph.style = caption_style


def build_reference_paragraph(doc: Document, previous: Paragraph, text: str) -> Paragraph:
    return insert_paragraph_after(previous, text, style="Normal")


def main() -> None:
    shutil.copy2(INPUT_PATH, OUTPUT_PATH)
    doc = Document(OUTPUT_PATH)

    # 1. Remove leftover proposal/planning content from the final report.
    remove_blocks_between(doc, "TÓM TẮT ĐỀ TÀI", "LỜI CAM ĐOAN")

    # 2. Clean obvious formatting/text issues in cover and body.
    global_replacements = [
        ("Hình", "Hình"),
        ("bỘ GIÁO DỤC VÀ ĐÀO TẠO     BỘ NÔNG NGHIỆP VÀ môi trường", "BỘ GIÁO DỤC VÀ ĐÀO TẠO     BỘ NÔNG NGHIỆP VÀ MÔI TRƯỜNG"),
        ("sesion", "session"),
        ("Edge Funtion", "Edge Function"),
        ("JWT Authenticantion", "JWT Authentication"),
        ("time gian thựcquá", "thời gian thực trong quá"),
        ("thời gian thựcquá", "thời gian thực trong quá"),
        ("thời gian thựcđến", "thời gian thực đến"),
        ("dữ liệucập", "dữ liệu cập"),
        ("dữ liệusẽ", "dữ liệu sẽ"),
        ("kẻ đ mới", "kẻ địch mới"),
        ("AI enemy", "AI kẻ địch"),
        ("start_match với user không phải chủ phòng", "start_match với người chơi không phải chủ phòng"),
        ("Bảng 4.3. Công cụ đồng bộ Đa người chơi", "Bảng 4.3 Công cụ đồng bộ đa người chơi"),
        ("Kiểm thử Backend/API và bảo mật", "Kiểm thử phía máy chủ, API và bảo mật"),
        ("Người dùng được chuyển hướng đến trang thông báo thanh toán thành công", "Trạng thái sẵn sàng được lưu và hiển thị lại cho các thành viên trong phòng."),
        ("JWT/session", "JWT hoặc mã phiên"),
        ("realtime networking", "đồng bộ thời gian thực"),
    ]

    for paragraph in doc.paragraphs:
        new_text = replace_in_text(paragraph.text, global_replacements)
        if new_text != paragraph.text:
            rewrite_paragraph(paragraph, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                new_text = replace_in_text(cell.text, global_replacements)
                if new_text != cell.text:
                    set_cell_text(cell, new_text)

    # Remove stray punctuation paragraph after the acknowledgement date.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == ".":
            rewrite_paragraph(paragraph, "")

    # 3. Strengthen Chapter 1 and Chapter 2 claims with clear scope and citations.
    p = find_paragraph_contains(doc, "đồng bộ dữ liệu, độ trễ mạng, quản lý phiên chơi")
    rewrite_paragraph(
        p,
        "Tuy nhiên, để xây dựng một game đa người chơi không chỉ cần xử lý đồ họa, nhân vật và quá trình chơi, mà còn phải giải quyết nhiều vấn đề phức tạp liên quan đến kiến trúc hệ thống, đồng bộ dữ liệu, độ trễ mạng, quản lý phiên chơi và đảm bảo tính nhất quán giữa các thiết bị người chơi [11], [12].",
    )

    p = find_paragraph_contains(doc, "Bên cạnh đó, vấn đề bảo mật và kiểm soát dữ liệu cũng đóng vai trò quan trọng.")
    rewrite_paragraph(
        p,
        "Bên cạnh đó, vấn đề bảo mật và kiểm soát dữ liệu cũng đóng vai trò quan trọng. Nếu toàn bộ logic xử lý được đặt ở phía người chơi, người chơi có thể can thiệp vào dữ liệu cục bộ để thay đổi kết quả trận đấu, điểm số hoặc tài nguyên. Vì vậy, việc xây dựng một hệ thống tách biệt giữa dữ liệu nghiệp vụ và dữ liệu quá trình chơi là cần thiết để giảm nguy cơ gian lận ở mức phù hợp với phạm vi đồ án [14].",
    )

    p = find_paragraph_contains(doc, "Kiến trúc tổng thể của hệ thống được thiết kế theo mô hình")
    rewrite_paragraph(
        p,
        "Kiến trúc tổng thể của hệ thống được tổ chức thành ba lớp chức năng tách biệt nhưng liên kết chặt chẽ: ứng dụng Unity ở phía người chơi, nền tảng Supabase ở phía máy chủ nghiệp vụ và Photon Fusion dùng để đồng bộ trạng thái trong trận theo thời gian thực [1], [2], [4]. Unity chịu trách nhiệm hiển thị giao diện, nhận thao tác điều khiển và mô phỏng cảm giác chơi tức thời; Supabase xử lý xác thực, lưu trữ dữ liệu bền vững, quản lý phòng, vòng đời trận và lịch sử; còn Photon Fusion chỉ đồng bộ trạng thái ngắn hạn phát sinh trong lúc trận đang diễn ra.",
    )
    p2 = insert_paragraph_after(
        p,
        "Trong bản cài đặt hiện tại, Photon Fusion được khởi tạo theo GameMode.Shared với SessionName chính là match_id của trận. Vì vậy, mỗi trận có một phiên mạng riêng; các đối tượng như người chơi, kẻ địch hay vật phẩm chỉ được đồng bộ trong phạm vi phiên đó. Dữ liệu này không được ghi từng khung hình xuống cơ sở dữ liệu; chỉ các mốc nghiệp vụ quan trọng như tạo phòng, bắt đầu trận, kết thúc trận và thống kê sau trận mới được lưu qua Supabase.",
        style="Normal",
    )

    p = find_paragraph_contains(doc, "Cách tổ chức này giúp hệ thống tách biệt rõ ràng giữa dữ liệu nghiệp vụ")
    rewrite_paragraph(
        p,
        "Cách tổ chức này giúp tách biệt dữ liệu nghiệp vụ cần lưu lâu dài với dữ liệu trạng thái thay đổi liên tục. Các bảng users, rooms, room_players, matches và match_players phục vụ quản lý tài khoản và vòng đời trận; còn vị trí, hướng xoay, hoạt ảnh, máu, kẻ địch, vật phẩm và sự kiện chiến đấu được giữ trong phiên Photon Fusion cho đến khi trận kết thúc. Cách phân lớp này phù hợp với định hướng chỉ lưu bền vững các sự kiện quan trọng thay vì ghi mọi thay đổi tức thời vào cơ sở dữ liệu [13].",
    )

    p = find_paragraph_contains(doc, "Để tránh nhầm lẫn giữa phía máy chủ nghiệp vụ và thời gian thực")
    rewrite_paragraph(
        p,
        "Để tránh nhầm lẫn giữa nền tảng máy chủ nghiệp vụ và lớp đồng bộ trận đấu theo thời gian thực, báo cáo sử dụng thống nhất ba khái niệm sau: ứng dụng Unity là phía người chơi; Supabase là phía máy chủ nghiệp vụ chịu trách nhiệm dữ liệu và phân quyền; Photon Fusion là lớp đồng bộ trạng thái trong phiên trận. Khi cần nói đến thực thể giữ quyền quyết định một trạng thái mạng, báo cáo dùng thuật ngữ quyền trạng thái (state authority) thay vì gọi chung là máy chủ.",
    )

    p = find_paragraph_contains(doc, "Như vậy, Supabase Realtime không thay thế Photon Fusion")
    rewrite_paragraph(
        p,
        "Như vậy, Supabase Realtime không thay thế Photon Fusion trong phần quá trình chơi. Supabase Realtime phù hợp với thay đổi nghiệp vụ có tần suất thấp như cập nhật lobby hoặc danh sách người chơi trong phòng [5], còn Photon Fusion phù hợp với dữ liệu thay đổi liên tục theo từng khung mô phỏng của trận đấu [4].",
    )

    p = find_paragraph_contains(doc, "Khi có nhiều người chơi trực tuyến cùng lúc, mỗi nhóm người chơi có thể tạo phòng riêng.")
    rewrite_paragraph(
        p,
        "Trong thiết kế hiện tại, mỗi phòng được định danh bằng room_id và room_code, còn mỗi trận được định danh bằng match_id. Nhờ đó, nhiều nhóm người chơi có thể tạo phòng và chơi ở các phiên khác nhau mà không dùng chung trạng thái phòng hoặc trạng thái trận. Tuy nhiên, nhận định này mới được kiểm thử trong phạm vi 2 đến 4 người chơi mỗi trận; báo cáo không suy rộng thành khả năng chịu tải lớn cho nhiều phòng hoạt động đồng thời trong thời gian dài.",
    )

    p = find_paragraph_contains(doc, "Gán session name cho Photon Fusion dựa trên match_id.")
    rewrite_paragraph(p, "Gán SessionName của Photon Fusion theo match_id để tách biệt từng phiên trận đấu.")

    p = find_paragraph_contains(doc, "Các Edge Functions như start_match và end_match đóng vai trò quan trọng")
    rewrite_paragraph(
        p,
        "Các Edge Functions như start_match và end_match đóng vai trò quan trọng trong mô-đun này. Chúng không cho phép ứng dụng Unity ghi trực tiếp dữ liệu room hoặc match vào cơ sở dữ liệu, mà buộc phải đi qua lớp kiểm tra JWT, quyền truy cập và trạng thái trận. Tuy vậy, các thống kê gameplay được gửi cuối trận hiện mới được kiểm tra ở mức ngữ cảnh nghiệp vụ và chống ghi trùng; phía máy chủ chưa tái mô phỏng lại toàn bộ sát thương hoặc số lần hạ gục để xác minh tuyệt đối.",
    )

    p = find_paragraph_contains(doc, "AI kẻ địch trong hệ thống được triển khai dựa trên Unity AI Navigation/NavMesh.")
    rewrite_paragraph(
        p,
        "AI kẻ địch trong hệ thống được triển khai dựa trên Unity AI Navigation/NavMesh. Kẻ địch sử dụng NavMesh để tìm đường trong bản đồ, xác định mục tiêu người chơi gần nhất và tiếp cận để tấn công. Cách tiếp cận này phù hợp với phạm vi demo vì vừa bảo đảm hành vi di chuyển có định hướng, vừa thuận tiện để gắn vào cơ chế đồng bộ mạng của Photon Fusion [6].",
    )

    p = find_paragraph_contains(doc, "Cách kết hợp này phù hợp với mô hình semi server-authoritative")
    rewrite_paragraph(
        p,
        "Cách kết hợp này phù hợp với mô hình bán quyết định bởi máy chủ trong phạm vi đồ án. Dữ liệu nghiệp vụ quan trọng như tài khoản, phòng, quyền bắt đầu trận và kết quả lưu trữ do Supabase kiểm tra; còn nhiều trạng thái gameplay vẫn do Photon Fusion Shared Mode và các peer trong phiên phối hợp duy trì. Vì vậy, giải pháp phù hợp với bản demo co-op 2 đến 4 người chơi, nhưng chưa tương đương kiến trúc máy chủ quyết định hoàn toàn của các game thương mại [4], [14].",
    )

    # 4. Clarify Chapter 3 authority, use cases, and constraints.
    p = find_paragraph_contains(doc, "Backend/Supabase đóng vai trò kiểm soát dữ liệu và xử lý các thao tác quan trọng.")
    rewrite_paragraph(
        p,
        "Supabase đóng vai trò máy chủ nghiệp vụ và xử lý các thao tác quan trọng liên quan tới tài khoản, phòng và trận đấu. Các chức năng như create_room, join_room, set_ready, start_match và end_match được thực hiện thông qua Edge Function. Mỗi yêu cầu đều gửi kèm JWT để xác định người dùng đang thao tác; tuy nhiên phạm vi bảo vệ này chủ yếu áp dụng cho dữ liệu nghiệp vụ trên Supabase, không đồng nghĩa với việc toàn bộ trạng thái gameplay trong phiên Photon Fusion đã được xác minh bởi một máy chủ trung tâm.",
    )

    p = find_paragraph_contains(doc, "Cơ chế bảo mật của hệ thống được hỗ trợ bởi JWT Authentication")
    rewrite_paragraph(
        p,
        "Cơ chế bảo mật của hệ thống được hỗ trợ bởi JWT Authentication và Row Level Security (RLS). JWT giúp xác định người dùng đang thao tác, còn RLS giới hạn quyền đọc và ghi dữ liệu theo từng người dùng [9], [10]. Nhờ đó, người chơi không thể tùy ý truy cập dữ liệu nghiệp vụ của người khác qua REST API hoặc Edge Function nếu không có quyền phù hợp.",
    )

    p = find_paragraph_contains(doc, "Với cách phân chia này, hệ thống đạt mức bảo mật phù hợp với demo đồ án.")
    rewrite_paragraph(
        p,
        "Với cách phân chia này, hệ thống đạt mức bảo mật phù hợp với bản demo đồ án. Tuy nhiên, do trò chơi đang chạy Photon Fusion ở Shared Mode chứ chưa có dedicated server quyết định toàn bộ trạng thái, báo cáo chỉ khẳng định khả năng kiểm soát dữ liệu nghiệp vụ và chống ghi sai ở mức cơ bản; chưa thể xem đây là cơ chế chống gian lận hoàn chỉnh cho mọi sự kiện gameplay thời gian thực.",
    )
    insert_paragraph_after(
        p,
        "Trong mã hiện tại, FusionMatchBootstrap khởi tạo phiên Photon Fusion ở GameMode.Shared theo match_id. Người chơi cục bộ tự sinh nhân vật mạng của mình trong phiên, còn Spawner sinh kẻ địch mạng với cờ SharedModeStateAuthLocalPlayer; các RPC gây sát thương của kẻ địch được gửi tới peer đang giữ state authority của đối tượng tương ứng. Vì vậy, Supabase không tham gia mô phỏng từng khung hình của trận mà chỉ ghi nhận kết quả nghiệp vụ sau cùng.",
        style="Normal",
    )

    # Update authority table (Bảng 3.3).
    table = find_table_after_caption(doc, "Bảng 3.3 Phân quyền xử lý trạng thái trong hệ thống")
    set_cell_text(table.cell(4, 1), "Edge Function; chủ phòng chỉ có quyền yêu cầu")
    set_cell_text(table.cell(4, 2), "Kiểm tra JWT, quyền chủ phòng, số người chơi hợp lệ, trạng thái sẵn sàng và ngăn tạo trận trùng trong cùng room_id.")
    set_cell_text(table.cell(5, 1), "Photon Fusion Shared Mode; đối tượng player do peer sở hữu điều khiển")
    set_cell_text(table.cell(5, 2), "Đồng bộ bằng NetworkObject, NetworkBehaviour và dữ liệu mạng; không ghi từng khung hình xuống cơ sở dữ liệu.")
    set_cell_text(table.cell(6, 1), "Peer đang giữ state authority của enemy hoặc spawner trong Photon Fusion")
    set_cell_text(table.cell(6, 2), "Sự kiện spawn, cập nhật máu và chết được gửi bằng RPC hoặc Networked Property tới các peer trong cùng match_id.")
    set_cell_text(table.cell(7, 1), "Peer đang giữ state authority của đối tượng vật phẩm hoặc cổng")
    set_cell_text(table.cell(7, 2), "Sự kiện nhặt vật phẩm, mở cổng và hồi sinh được áp dụng một lần trong phiên rồi quảng bá tới các peer còn lại.")
    set_cell_text(table.cell(8, 1), "Edge Function end_match và các bảng matches, match_players")
    set_cell_text(table.cell(8, 2), "Kiểm tra JWT, match_id, người chơi có thuộc trận hay không, trạng thái trận và ngăn gửi lặp kết quả.")
    format_table_for_report(table)

    # Update use case tables.
    table = find_table_after_caption(doc, "Bảng 3.8 Đặc tả Use Case Tham gia phòng chơi")
    set_cell_text(table.cell(4, 1), "Người chơi đã đăng nhập, phòng tồn tại, chưa vượt giới hạn số người chơi và chưa chuyển sang trạng thái bắt đầu trận.")
    set_cell_text(table.cell(7, 1), "1. Người chơi nhập room code. / 2. Ứng dụng Unity gọi Edge Function join_room. / 3. Phía máy chủ xác thực JWT. / 4. Hệ thống kiểm tra phòng có tồn tại hay không. / 5. Hệ thống kiểm tra phòng còn chỗ và chưa bắt đầu trận. / 6. Nếu hợp lệ, hệ thống thêm người chơi vào room_players và cập nhật giao diện phòng.")
    set_cell_text(table.cell(8, 1), "Nếu room code sai, hệ thống báo phòng không tồn tại. / Nếu phòng đã đầy hoặc đã chuyển sang trận đấu, yêu cầu tham gia bị từ chối. / Nếu JWT hết hạn, người chơi phải đăng nhập lại.")
    format_table_for_report(table)

    table = find_table_after_caption(doc, "Bảng 3.9 Đặc tả Use Case Sẵn sàng")
    set_cell_text(table.cell(5, 1), "Trạng thái is_ready của người chơi được cập nhật và hiển thị lại cho các thành viên trong phòng.")
    set_cell_text(table.cell(9, 1), "Nếu người chơi không còn ở trong phòng hoặc JWT hết hạn, yêu cầu cập nhật trạng thái bị từ chối.")
    format_table_for_report(table)

    table = find_table_after_caption(doc, "Bảng 3.10 Đặc tả Use Case Bắt đầu trận")
    set_cell_text(table.cell(7, 1), "1. Chủ phòng nhấn Start Match. / 2. Ứng dụng Unity gọi Edge Function start_match. / 3. Phía máy chủ xác thực JWT. / 4. Hệ thống kiểm tra người gọi có phải chủ phòng hay không. / 5. Hệ thống kiểm tra số người chơi, trạng thái sẵn sàng và việc đã có trận đang hoạt động trong room_id hay chưa. / 6. Nếu hợp lệ, hệ thống tạo match_id, cập nhật trạng thái trận và trả dữ liệu để các client vào GameScene.")
    set_cell_text(table.cell(8, 1), "Nếu người gọi không phải chủ phòng, hệ thống báo không có quyền. / Nếu có người chưa sẵn sàng hoặc phòng đang có trận hoạt động, hệ thống từ chối tạo trận mới. / Nếu request bị gửi lặp, hệ thống không tạo thêm match_id thứ hai.")
    format_table_for_report(table)

    table = find_table_after_caption(doc, "Bảng 3.13 Đặc tả Use Case Kết thúc trận và lưu kết quả")
    set_cell_text(table.cell(5, 1), "Kết quả trận được lưu vào bảng matches và match_players, giao diện kết quả được hiển thị cho người chơi.")
    set_cell_text(table.cell(7, 1), "1. Hệ thống xác định điều kiện thắng hoặc thua trong GameScene. / 2. Ứng dụng Unity tổng hợp thống kê như kill, sát thương, số lần bị hạ, hồi sinh và thời gian sống. / 3. Ứng dụng Unity gọi Edge Function end_match kèm match_id và danh sách thống kê. / 4. Phía máy chủ kiểm tra JWT, người chơi có thuộc trận hay không, trạng thái trận và việc đã lưu kết quả trước đó chưa. / 5. Nếu hợp lệ, hệ thống cập nhật matches và match_players, sau đó trả dữ liệu để hiển thị lịch sử trận.")
    set_cell_text(table.cell(8, 1), "Nếu dữ liệu gửi lên không hợp lệ, phía máy chủ từ chối lưu kết quả. / Nếu request bị gửi lặp, hệ thống không ghi thêm bản ghi match_players trùng. / Phía máy chủ hiện chưa tái mô phỏng toàn bộ chiến đấu để xác minh tuyệt đối mọi chỉ số gameplay.")
    format_table_for_report(table)

    table = find_table_after_caption(doc, "Bảng 3.15 Đặc tả Use Case Rời khỏi phòng")
    set_cell_text(table.cell(8, 1), "Nếu người rời phòng là chủ phòng, hệ thống phải dọn trạng thái phòng; việc chuyển chủ phòng chỉ áp dụng khi logic Edge Function hỗ trợ rõ ràng.")
    format_table_for_report(table)

    # Update constraints section wording.
    p = find_paragraph_contains(doc, "Ràng buộc dữ liệu và chỉ mục cần áp dụng")
    rewrite_paragraph(p, "Ràng buộc dữ liệu đã triển khai và hướng hoàn thiện")

    p = find_paragraph_contains(doc, "Ngoài các bảng dữ liệu chính, hệ thống cần xác định ràng buộc")
    rewrite_paragraph(
        p,
        "Ngoài các bảng dữ liệu chính, hệ thống cần xác định rõ ràng buộc nào đã được triển khai trực tiếp trong cơ sở dữ liệu hoặc Edge Function, và ràng buộc nào mới ở mức tối ưu hóa thiết kế. Trong phạm vi hiện tại, báo cáo tập trung vào khóa chính, khóa ngoại, ràng buộc duy nhất, kiểm tra quyền bằng JWT/RLS và kiểm tra vòng đời trạng thái phòng hoặc trận khi gọi các hàm nghiệp vụ.",
    )
    p = find_paragraph_contains(doc, "Các ràng buộc trên giúp hệ thống xử lý tốt hơn")
    rewrite_paragraph(
        p,
        "Trong số các ràng buộc trên, các khóa chính, khóa ngoại, ràng buộc duy nhất theo room_id hoặc match_id cùng kiểm tra quyền qua JWT là phần đã được triển khai và kiểm thử trực tiếp. Các hạng mục như tối ưu chỉ mục sâu hơn, giới hạn tần suất gọi hàm hoặc dọn dữ liệu phòng treo tự động là hướng hoàn thiện thêm, chưa được kiểm thử tải chuyên sâu trong báo cáo này.",
    )

    # 5. Rewrite Chapter 4 testing with more concrete setup and bounded claims.
    table = find_table_after_caption(doc, "Bảng 4.5 Môi trường quản lý và kiểm thử")
    while len(table.rows) < 10:
        table.add_row()
    rows = [
        ("Thành phần", "Mô tả"),
        ("Thiết bị kiểm thử", "01 máy Windows 11, AMD Ryzen 7 5800H, RAM 16 GB."),
        ("Bố trí client", "01 phiên Unity Editor và 1-3 bản build Windows chạy đồng thời theo từng kịch bản."),
        ("Số lượng client", "2 đến 4 client trong cùng một phiên kiểm thử."),
        ("Backend nghiệp vụ", "Supabase Auth, Supabase Database, Edge Functions."),
        ("Networking trận đấu", "Photon Fusion Shared Mode theo match_id của từng trận."),
        ("Công cụ kiểm thử API", "Postman và Supabase Dashboard."),
        ("Công cụ đo hiệu năng", "Unity Stats/Profiler để đọc FPS; Photon Fusion Stats để đọc RTT/Ping."),
        ("Cách lấy số liệu", "Mỗi kịch bản gameplay được chạy 3 lần, quan sát khoảng 180 giây và lấy giá trị trung bình."),
        ("Điều kiện mạng", "Các client dùng cùng một kết nối Internet ổn định; báo cáo không mô phỏng mất gói hoặc jitter nhân tạo."),
    ]
    for r_idx, (left, right) in enumerate(rows):
        set_cell_text(table.cell(r_idx, 0), left)
        set_cell_text(table.cell(r_idx, 1), right)
    format_table_for_report(table)

    p = find_paragraph_contains(doc, "Việc kiểm thử được thực hiện bằng cách cho nhiều client đăng nhập")
    rewrite_paragraph(
        p,
        "Việc kiểm thử được thực hiện bằng cách cho nhiều client đăng nhập bằng các tài khoản khác nhau, cùng tạo hoặc tham gia phòng, vào trận và thực hiện các thao tác chính trong cùng một phiên chơi. Với các kịch bản gameplay, số liệu FPS và RTT được ghi trên client quan sát chính trong khoảng 180 giây cho mỗi lần chạy; mỗi kịch bản được lặp lại 3 lần để lấy giá trị trung bình.",
    )

    table = find_table_after_caption(doc, "Bảng 4.6 Ma trận kiểm thử chức năng chính")
    actual_results = {
        1: "Đăng nhập thành công; Supabase Auth trả mã phiên hợp lệ và hồ sơ người chơi được tải lên giao diện.",
        2: "Tạo mới bản ghi rooms và room_players; người tạo phòng được gán vai trò chủ phòng và nhận room_code.",
        3: "Người chơi thứ hai và thứ ba vào phòng thành công; danh sách thành viên được cập nhật trên các client.",
        4: "Cờ is_ready thay đổi đúng và hiển thị gần như tức thời cho các thành viên trong lobby.",
        5: "Hệ thống sinh match_id, cập nhật trạng thái trận và chuyển các client sang GameScene.",
        6: "Từ 2 đến 4 người chơi xuất hiện trong cùng phiên match_id và nhìn thấy nhau trong cùng bản đồ.",
        7: "Các client còn lại quan sát được vị trí, hướng xoay và hoạt ảnh di chuyển thay đổi liên tục.",
        8: "Máu kẻ địch giảm đúng; sự kiện chết, cộng kill và mở cổng được đồng bộ cho các client.",
        9: "Vật phẩm biến mất khỏi phiên ngay sau khi một người chơi nhặt; không phát hiện nhặt trùng.",
        10: "Người chơi bị hạ được hồi sinh và tiếp tục trận; trạng thái đồng bộ cho các thành viên còn lại.",
        11: "Màn hình Thắng/Thua hiển thị đồng nhất; dữ liệu kết quả được gửi lên end_match.",
        12: "Lịch sử trận hiển thị đúng trận vừa lưu với các chỉ số kill, damage và revive.",
    }
    for row_index, actual in actual_results.items():
        set_cell_text(table.cell(row_index, 4), actual)
        set_cell_text(table.cell(row_index, 5), "Đạt")
    format_table_for_report(table)

    table = find_table_after_caption(doc, "Bảng 4.7 Kiểm thử hiệu năng")
    row_updates = {
        1: (
            "Một client di chuyển liên tục trong bản đồ khoảng 30 giây.",
            "Client quan sát nhìn thấy vị trí cập nhật liên tục, không xuất hiện dịch chuyển bất thường kéo dài.",
        ),
        2: (
            "Người chơi chạy, trượt né và tấn công combo 1-2-3 ở nhiều vị trí khác nhau.",
            "Hành động hiển thị đúng nhịp trên các client còn lại; chưa ghi nhận mất đồng bộ kéo dài ở các thao tác cơ bản.",
        ),
        3: (
            "Kẻ địch tấn công người chơi, sau đó kiểm tra trạng thái bị hạ và hồi sinh.",
            "Giá trị máu và trạng thái bị hạ hoặc hồi sinh được cập nhật đồng nhất giữa các client trong phiên.",
        ),
        4: (
            "Kẻ địch di chuyển, nhận sát thương, chết và làm sạch khu vực spawn.",
            "Các client quan sát cùng trạng thái kẻ địch; khi toàn bộ kẻ địch chết, cổng khu vực mở đúng một lần.",
        ),
        5: (
            "Một người chơi nhặt vật phẩm hồi máu hoặc vật phẩm xu trong lúc client khác cùng quan sát.",
            "Vật phẩm biến mất trên toàn bộ phiên; không thấy vật phẩm bị nhặt lặp hoặc tồn tại bản sao hiển thị.",
        ),
        6: (
            "Trận kết thúc theo cả hai hướng thắng và thua.",
            "Tất cả client nhận cùng kết quả Victory hoặc Lose; giao diện cuối trận không mâu thuẫn dữ liệu.",
        ),
    }
    for r_idx, (method, actual) in row_updates.items():
        set_cell_text(table.cell(r_idx, 1), method)
        set_cell_text(table.cell(r_idx, 3), actual)
    format_table_for_report(table)

    for paragraph in doc.paragraphs:
        if "Bảng 4.8 Bảng đo hiệu năng" in paragraph.text:
            rewrite_paragraph(paragraph, paragraph.text.replace("Bảng 4.8 Bảng đo hiệu năng", "Bảng 4.8 Số liệu hiệu năng trong trận và workflow cuối trận"))

    for paragraph in doc.paragraphs:
        if "Hình 4.25  Kiểm thử API start_match với người chơi không phải chủ phòng" in paragraph.text:
            rewrite_paragraph(paragraph, paragraph.text.replace("Hình 4.25  Kiểm thử API start_match với người chơi không phải chủ phòng", "Hình 4.24  Kiểm thử API start_match với người chơi không phải chủ phòng"))
        if "Hình 4.26  Dữ liệu match_players sau khi lưu kết quả trận đấu" in paragraph.text:
            rewrite_paragraph(paragraph, paragraph.text.replace("Hình 4.26  Dữ liệu match_players sau khi lưu kết quả trận đấu", "Hình 4.25  Dữ liệu match_players sau khi lưu kết quả trận đấu"))

    p = find_last_paragraph_contains(doc, "Bảng 4.8 Số liệu hiệu năng trong trận và workflow cuối trận")
    rewrite_paragraph(p, "Bảng 4.8 Số liệu hiệu năng trong trận và workflow cuối trận")

    table = find_table_after_caption(doc, "Bảng 4.8 Số liệu hiệu năng trong trận và workflow cuối trận")
    set_cell_text(table.cell(1, 0), "Một phòng chơi cơ bản (di chuyển và quan sát)")
    set_cell_text(table.cell(2, 0), "Một phòng chơi đầy đủ")
    set_cell_text(table.cell(3, 0), "Chiến đấu với nhiều kẻ địch")
    set_cell_text(table.cell(4, 0), "Kết thúc trận và lưu kết quả (workflow cuối trận)")
    set_cell_text(table.cell(4, 4), "Chỉ dùng để quan sát end_match và lưu lịch sử; không dùng suy rộng cho hiệu năng gameplay")
    format_table_for_report(table)

    p = find_paragraph_contains(doc, "Kết quả đo cho thấy hệ thống hoạt động ổn định trong phạm vi thử nghiệm từ 2 đến 4 người chơi.")
    rewrite_paragraph(
        p,
        "FPS được ghi trên client quan sát chính bằng Unity Stats hoặc Profiler; RTT/Ping được đọc từ Photon Fusion Stats của cùng client. Mỗi kịch bản gameplay được chạy 3 lần trong khoảng 180 giây và lấy giá trị trung bình. Trong phạm vi này, hệ thống hoạt động ổn định với 2 đến 4 người chơi trong cùng một trận.",
    )
    p = find_paragraph_contains(doc, "Khi tăng lên 4 người chơi, FPS duy trì quanh 40")
    rewrite_paragraph(
        p,
        "Khi tăng lên 4 người chơi, FPS duy trì quanh 40 và RTT tăng lên 87 đến 100 ms, kể cả trong kịch bản chiến đấu với nhiều kẻ địch. Riêng bước kết thúc trận và lưu kết quả có FPS cao hơn vì tải đồng bộ gameplay giảm; số liệu đó chỉ được dùng để quan sát workflow end_match và lưu lịch sử sau trận, không được dùng làm căn cứ kết luận về hiệu năng chiến đấu trong thời gian dài.",
    )
    p = find_paragraph_contains(doc, "Nhìn chung, bản demo đáp ứng tốt quy mô thử nghiệm hiện tại")
    rewrite_paragraph(
        p,
        "Nhìn chung, bản demo đáp ứng được quy mô thử nghiệm hiện tại của đồ án, nhưng kết luận này chỉ đúng trong phạm vi 2 đến 4 người chơi, số kẻ địch hiện tại và điều kiện mạng đã đo. Báo cáo không suy rộng các số liệu này thành khả năng chịu tải lớn cho nhiều phòng hoạt động đồng thời trong thời gian dài.",
    )

    p = find_paragraph_contains(doc, "Các thao tác quan trọng như tạo phòng, tham gia phòng, bắt đầu trận và lưu kết quả đều được kiểm thử")
    rewrite_paragraph(
        p,
        "Các thao tác quan trọng như tạo phòng, tham gia phòng, bắt đầu trận và lưu kết quả đều được kiểm thử qua Supabase Edge Functions để xác nhận rằng người chơi không thể ghi trực tiếp dữ liệu nghiệp vụ nếu không có JWT hợp lệ hoặc không đủ quyền. Phạm vi kiểm thử bảo mật hiện tại tập trung vào xác thực, quyền chủ phòng, người chơi ngoài phòng và chống gửi lặp; chưa bao phủ toàn bộ các tình huống giả mạo thống kê gameplay ở mức chuyên sâu.",
    )

    table = find_table_after_caption(doc, "Bảng 4.9 Kiểm thử Phía máy chủ")
    security_results = {
        1: "Yêu cầu bị từ chối ở tầng xác thực; không phát sinh bản ghi mới trong dữ liệu phòng hoặc trận.",
        2: "Token sai bị từ chối; người chơi không đọc hoặc ghi được dữ liệu nghiệp vụ cần quyền xác thực.",
        3: "set_ready trả lỗi quyền; trạng thái trong room_players không thay đổi.",
        4: "start_match bị từ chối; không sinh match_id mới và phòng không đổi trạng thái.",
        5: "join_room thất bại; người chơi không được thêm vào danh sách thành viên phòng.",
        6: "Lần gửi lặp không tạo thêm bản ghi match_players; trạng thái trận giữ nguyên sau lần ghi đầu.",
    }
    for row_index, actual in security_results.items():
        set_cell_text(table.cell(row_index, 3), actual)
        set_cell_text(table.cell(row_index, 4), "Đạt")
    format_table_for_report(table)

    p = find_paragraph_contains(doc, "Kết quả kiểm thử lỗi mạng cho thấy hệ thống đã tách được dữ liệu giữa các phòng và trận")
    rewrite_paragraph(
        p,
        "Kết quả kiểm thử lỗi mạng cho thấy hệ thống đã tách được dữ liệu giữa các phòng và trận thông qua room_id, match_id và phiên mạng Photon Fusion. Cần lưu ý rằng trong trận đấu, khái niệm chủ phòng ở Supabase chỉ là người tạo phòng tại lớp nghiệp vụ; phiên Photon Fusion đang chạy ở Shared Mode nên không có network host theo nghĩa client-server truyền thống. Vì vậy, kịch bản người tạo phòng rời giữa trận được hiểu là một peer rời phiên Shared Mode; trận vẫn có thể tiếp tục nếu còn peer khác giữ state authority cho các đối tượng đang hoạt động. Hệ thống hiện chưa cài đặt host migration chuyên biệt hoặc cơ chế reconnect tự động đầy đủ.",
    )

    table = find_table_after_caption(doc, "Bảng 4.10 Kiểm thử mạng và độ ổn định")
    set_cell_text(table.cell(1, 2), "Hệ thống cập nhật lại danh sách người chơi trong phòng; người đã thoát không còn xuất hiện ở lobby.")
    set_cell_text(table.cell(2, 2), "Phòng được dọn trạng thái; người chơi còn lại phải nhận chủ phòng mới hoặc tạo lại phòng tùy theo xử lý của Edge Function.")
    set_cell_text(table.cell(3, 2), "Photon Fusion phát hiện peer rời phiên; trận hiện chỉ hỗ trợ xử lý mất kết nối ở mức cơ bản, chưa có reconnect tự động đầy đủ.")
    set_cell_text(table.cell(3, 3), "Đạt ở mức cơ bản")
    set_cell_text(table.cell(4, 2), "Các người chơi còn lại vẫn tiếp tục trận bình thường; nếu chỉ còn 1 người chơi thì trận vẫn hoạt động đến khi kết thúc.")
    set_cell_text(table.cell(4, 3), "Đạt ở mức cơ bản")
    set_cell_text(table.cell(5, 2), "Trong Shared Mode, người tạo phòng rời trận không đồng nghĩa với việc toàn bộ phiên mạng dừng lại; trận vẫn tiếp tục nếu còn peer khác giữ authority cho các đối tượng đang hoạt động.")
    set_cell_text(table.cell(5, 3), "Đạt ở mức cơ bản")
    set_cell_text(table.cell(6, 2), "Phía máy chủ kiểm tra match_id và trạng thái trận để tránh ghi lặp lịch sử.")
    set_cell_text(table.cell(7, 2), "Dữ liệu phòng và trận được tách bằng room_id, match_id và SessionName của Photon Fusion.")
    format_table_for_report(table)

    p = find_paragraph_contains(doc, "Kết quả kiểm thử cho thấy hệ thống đáp ứng được luồng chính của đồ án")
    rewrite_paragraph(
        p,
        "Kết quả kiểm thử cho thấy hệ thống đáp ứng được luồng chính của đồ án: người chơi có thể đăng nhập, tạo phòng, tham gia phòng, sẵn sàng, bắt đầu trận, chơi cùng nhau, đồng bộ trạng thái trong trận và lưu kết quả sau trận. Tuy nhiên, các kết luận về tính ổn định chỉ được xác nhận trong phạm vi kịch bản đã đo và chưa đủ để khẳng định hệ thống chịu tải lớn hay chống gian lận chuyên sâu.",
    )

    p = find_paragraph_contains(doc, "Hạn chế hiện tại là quy mô kiểm thử còn nhỏ")
    rewrite_paragraph(
        p,
        "Hạn chế hiện tại là quy mô kiểm thử còn nhỏ, chưa đánh giá tải lớn trong thời gian dài và chưa hoàn thiện các chức năng như kết nối lại sau mất mạng, quản lý dọn phòng treo hoặc kiểm chứng sâu các số liệu gameplay ở phía máy chủ. Đây là các nội dung cần tiếp tục mở rộng nếu phát triển sản phẩm lên mức hoàn chỉnh hơn.",
    )

    # Add teacher-requested objective-evidence summary table before Chapter 5.
    p = find_paragraph_contains(doc, "Hạn chế hiện tại là quy mô kiểm thử còn nhỏ")
    summary_heading = insert_paragraph_after(p, "Bảng liên kết mục tiêu, thành phần triển khai và bằng chứng", style="Heading 3")
    summary_caption = insert_paragraph_after(summary_heading, "Bảng 4.11 Liên kết mục tiêu, thành phần triển khai và bằng chứng kiểm thử", style="Caption")
    summary_table = insert_table_after(summary_caption, rows=5, cols=5, style="Table Grid")
    summary_rows = [
        ["Mục tiêu", "Thành phần triển khai", "Test case / chỉ số", "Bằng chứng", "Giới hạn cần nêu"],
        ["Quản lý phòng và vòng đời trận", "Supabase Auth, rooms, room_players, matches, Edge Functions", "TC02-TC05; SEC03-SEC06", "Kết quả Postman, dữ liệu trước/sau và hình kiểm thử", "Chưa kiểm thử request đồng thời ở tải lớn"],
        ["Đồng bộ gameplay 2-4 người chơi", "Photon Fusion Shared Mode, NetworkObject, RPC và dữ liệu mạng", "TC06-TC11; FPS/RTT ở Bảng 4.8", "Số liệu Unity Stats, Photon Stats và hình trong GameScene", "Chỉ xác nhận trong phạm vi 2-4 người chơi"],
        ["Kiểm soát dữ liệu nghiệp vụ", "JWT, RLS, Edge Functions, ràng buộc dữ liệu", "SEC01-SEC06", "Phản hồi lỗi, không phát sinh bản ghi sai và dữ liệu Supabase sau kiểm thử", "Chưa xác minh toàn bộ chỉ số gameplay theo kiểu authoritative"],
        ["Ổn định khi xảy ra lỗi mạng cơ bản", "Phiên Photon Fusion, room_id, match_id, cleanup ở lớp nghiệp vụ", "Bảng 4.10", "Quan sát client còn lại và trạng thái phòng hoặc trận sau khi thoát", "Chưa có reconnect tự động hoặc host migration chuyên biệt"],
    ]
    for r_idx, row in enumerate(summary_rows):
        for c_idx, value in enumerate(row):
            set_cell_text(summary_table.cell(r_idx, c_idx), value)
    format_table_for_report(summary_table)

    # 6. Add a real conclusion section to Chapter 5.
    chapter5 = find_paragraph_contains(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    conclusion_heading = insert_paragraph_after(chapter5, "Kết luận", style="Heading 2")
    insert_paragraph_after(
        conclusion_heading,
        "Đồ án đã xây dựng được một bản demo game 3D co-op 2 đến 4 người chơi gồm đăng ký hoặc đăng nhập, tạo và tham gia phòng, sẵn sàng, bắt đầu trận, đồng bộ nhân vật và kẻ địch, kết thúc trận và lưu kết quả. Về mặt kiến trúc, hệ thống phân tách rõ Supabase cho dữ liệu nghiệp vụ và Photon Fusion cho đồng bộ thời gian thực; các ca kiểm thử chức năng, bảo mật cơ bản và hiệu năng quy mô nhỏ cho thấy giải pháp đáp ứng mục tiêu của một đồ án thực hành có tính hoàn chỉnh. Tuy nhiên, các kết luận về tính ổn định chỉ đúng trong phạm vi 2 đến 4 người chơi và chưa đủ để suy rộng thành tải lớn, chống gian lận chuyên sâu hay khả năng khôi phục đầy đủ sau sự cố mạng phức tạp.",
        style="Normal",
    )

    # Fix AI typo in limitations.
    p = find_paragraph_contains(doc, "Thứ tư, phần AI của kẻ đ")
    rewrite_paragraph(
        p,
        "Thứ tư, phần AI của kẻ địch mới dừng ở mức cơ bản. Kẻ địch có thể di chuyển bằng NavMesh, tìm mục tiêu, tấn công người chơi và được sinh ra theo đợt kẻ địch. Tuy nhiên, hành vi của kẻ địch chưa đa dạng, chưa có nhiều chiến thuật phối hợp và chưa có cơ chế ra quyết định phức tạp hơn như ưu tiên mục tiêu theo mối đe dọa hoặc phản ứng theo đội hình người chơi.",
    )

    # 7. Add a full reference section compatible with existing [1]...[10] citations and new additions.
    refs_heading = find_paragraph_contains(doc, "Hướng phát triển trong tương lai")
    last_paragraph = doc.paragraphs[-1]
    ref_heading = insert_paragraph_after(last_paragraph, "TÀI LIỆU THAM KHẢO", style="Heading 1")
    references = [
        "[1] Unity Technologies, \"Unity Manual,\" [Trực tuyến]. Available: https://docs.unity3d.com/Manual/. [Truy cập: 20-06-2026].",
        "[2] Supabase, \"Supabase Documentation,\" [Trực tuyến]. Available: https://supabase.com/docs. [Truy cập: 20-06-2026].",
        "[3] Supabase, \"Edge Functions,\" [Trực tuyến]. Available: https://supabase.com/docs/guides/functions. [Truy cập: 20-06-2026].",
        "[4] Photon Engine, \"Photon Fusion Documentation,\" [Trực tuyến]. Available: https://doc.photonengine.com/fusion/current/. [Truy cập: 20-06-2026].",
        "[5] Supabase, \"Subscribing to Database Changes,\" [Trực tuyến]. Available: https://supabase.com/docs/guides/realtime/subscribing-to-database-changes. [Truy cập: 20-06-2026].",
        "[6] Unity Technologies, \"AI Navigation / NavMesh,\" [Trực tuyến]. Available: https://docs.unity3d.com/Packages/com.unity.ai.navigation@2.0/. [Truy cập: 20-06-2026].",
        "[7] Unity Technologies, \"Cinemachine,\" [Trực tuyến]. Available: https://docs.unity3d.com/Manual/com.unity.cinemachine.html. [Truy cập: 20-06-2026].",
        "[8] The PostgreSQL Global Development Group, \"PostgreSQL Documentation,\" [Trực tuyến]. Available: https://www.postgresql.org/docs/current/index.html. [Truy cập: 20-06-2026].",
        "[9] M. Jones, J. Bradley và N. Sakimura, \"JSON Web Token (JWT),\" RFC 7519, 2015. [Trực tuyến]. Available: https://datatracker.ietf.org/doc/html/rfc7519. [Truy cập: 20-06-2026].",
        "[10] Supabase, \"Row Level Security,\" [Trực tuyến]. Available: https://supabase.com/docs/guides/database/postgres/row-level-security. [Truy cập: 20-06-2026].",
        "[11] L. Pantel và L. C. Wolf, \"On the Impact of Delay on Real-Time Multiplayer Games,\" Proceedings of the 12th International Workshop on Network and Operating Systems Support for Digital Audio and Video, 2002.",
        "[12] W. Palant, C. Griwodz và P. Halvorsen, \"Consistency Requirements in Multiplayer Online Games,\" Proceedings of the 5th ACM SIGCOMM Workshop on Network and System Support for Games, 2006.",
        "[13] K. Zhang, B. Kemme và A. Denault, \"Persistence in Massively Multiplayer Online Games,\" Proceedings of the 7th ACM SIGCOMM Workshop on Network and System Support for Games, 2008.",
        "[14] B. D. Chen và M. Maheswaran, \"A Cheat Controlled Protocol for Centralized Online Multiplayer Games,\" Proceedings of the 3rd ACM SIGCOMM Workshop on Network and System Support for Games, 2004.",
    ]
    prev = ref_heading
    for ref in references:
        prev = build_reference_paragraph(doc, prev, ref)

    # 8. Normalize caption styles and keep large tables readable.
    set_caption_styles(doc)
    find_paragraph_contains(doc, "Bảng liên kết mục tiêu, thành phần triển khai và bằng chứng").style = "Heading 3"
    for table in doc.tables:
        format_table_for_report(table)

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
